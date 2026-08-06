"""End-to-end incremental-indexing behavior against live Postgres + Qdrant:
unchanged re-run does nothing, an edited file re-embeds and bumps version,
and a deleted file's vectors/registry row are removed."""

import uuid

import pytest
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from app.config import get_settings
from app.core.retrieval.qdrant_store import get_qdrant_client
from app.ingestion.metadata_resolver import MetadataResolver
from app.ingestion.pipeline import ingest_file
from app.ingestion.registry_sync import sync_folder
from app.models.document import Document

pytestmark = pytest.mark.integration


@pytest.fixture
async def session():
    """Uses a per-test tenant_id so committed rows can never collide with
    real demo tenants (acme/globex) or leak into an unrelated ingestion
    run's deletion-detection scope, and cleans them up afterward since
    committed data survives the session rollback below."""
    settings = get_settings()
    engine = create_async_engine(settings.database_url)
    session_factory = async_sessionmaker(bind=engine, expire_on_commit=False)
    tenant_id = f"test-tenant-{uuid.uuid4().hex[:8]}"

    async with session_factory() as s:
        s.test_tenant_id = tenant_id
        yield s
        await s.rollback()
        await s.execute(delete(Document).where(Document.tenant_id == tenant_id))
        await s.commit()
    await engine.dispose()


def _make_corpus(tmp_path, tenant_id: str):
    root = tmp_path / f"corpus-{uuid.uuid4().hex[:8]}"
    tenant_dir = root / tenant_id / "engineering" / "policies"
    tenant_dir.mkdir(parents=True)
    (tenant_dir / "policy.docx").write_bytes(_tiny_docx("Version one content."))
    return root, tenant_dir / "policy.docx"


def _tiny_docx(body_text: str) -> bytes:
    import io

    import docx

    doc = docx.Document()
    doc.add_heading("Policy", level=1)
    doc.add_paragraph(body_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _points_for_document(collection_name: str, document_id: str) -> int:
    from qdrant_client.http import models as qmodels

    client = get_qdrant_client()
    if not client.collection_exists(collection_name):
        return 0
    result, _ = client.scroll(
        collection_name=collection_name,
        scroll_filter=qmodels.Filter(
            must=[qmodels.FieldCondition(key="document_id", match=qmodels.MatchValue(value=document_id))]
        ),
        limit=1000,
    )
    return len(result)


async def test_unchanged_rerun_does_not_reembed(session, tmp_path):
    root, file_path = _make_corpus(tmp_path, session.test_tenant_id)
    resolver = MetadataResolver(root)

    first = await ingest_file(session, root, file_path, resolver.resolve(file_path))
    await session.commit()
    assert first.status == "ingested"
    assert first.num_leaf_chunks > 0

    second = await ingest_file(session, root, file_path, resolver.resolve(file_path))
    await session.commit()

    assert second.status == "unchanged"
    assert second.document_id == first.document_id


async def test_edited_file_reembeds_and_bumps_version(session, tmp_path):
    root, file_path = _make_corpus(tmp_path, session.test_tenant_id)
    resolver = MetadataResolver(root)

    first = await ingest_file(session, root, file_path, resolver.resolve(file_path))
    await session.commit()

    file_path.write_bytes(_tiny_docx("Version two, substantially different content now."))

    second = await ingest_file(session, root, file_path, resolver.resolve(file_path))
    await session.commit()

    assert second.status == "ingested"
    assert second.document_id == first.document_id

    result = await session.execute(select(Document).where(Document.id == uuid.UUID(second.document_id)))
    document = result.scalar_one()
    assert document.current_version == 2

    settings = get_settings()
    remaining = _points_for_document(settings.embedding_collection_name, second.document_id)
    assert remaining == second.num_leaf_chunks


async def test_deleted_file_removed_via_sync_folder(session, tmp_path):
    root, file_path = _make_corpus(tmp_path, session.test_tenant_id)
    expected_source_path = f"{session.test_tenant_id}/engineering/policies/policy.docx"

    summary1 = await sync_folder(session, root)
    assert len(summary1.ingested) == 1

    result = await session.execute(
        select(Document).where(Document.source_path == expected_source_path)
    )
    document = result.scalar_one()
    document_id = str(document.id)

    file_path.unlink()

    summary2 = await sync_folder(session, root)
    assert expected_source_path in summary2.deleted

    await session.refresh(document)
    assert document.status == "deleted"

    settings = get_settings()
    remaining = _points_for_document(settings.embedding_collection_name, document_id)
    assert remaining == 0
