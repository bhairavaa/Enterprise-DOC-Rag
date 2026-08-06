from pathlib import Path

import docx
from llama_index.core import Document


def load_docx(file_path: Path) -> list[Document]:
    """One Document per top-level section, split on paragraphs whose style
    is a 'Heading N' style. DOCX has no page concept (it's reflow-based), so
    section_title is the only structural metadata; the document falls back
    to a single unsectioned Document if it has no headings at all."""
    doc = docx.Document(str(file_path))

    documents: list[Document] = []
    current_title: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        text = "\n".join(current_lines).strip()
        if text:
            documents.append(Document(text=text, metadata={"section_title": current_title}))

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading") or style_name == "Title":
            flush()
            current_title = text
            current_lines = []
        else:
            current_lines.append(text)

    flush()
    return documents
